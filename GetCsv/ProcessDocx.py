from docx.api import Document
import os

#  處理 Docx 的表格資料，讓表格資料以正確的順序被讀取

def ProcessDocx(file_path, test=False):
    document = Document(file_path)

    for table in document.tables:

        table_contents = ""

        keys = None
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)

            if i == 0:
                keys = tuple(text)
                continue
            row_data = dict(zip(keys, text))
            row_text = str(row_data)
            table_contents += (row_text + '\n')

        for ch in ['\\n', '\\u3000', '\'', '{', '}']:
            table_contents = table_contents.replace(ch,'')

        new_para = document.add_paragraph(table_contents)
        table._element.addnext(new_para._p)

        table._element.getparent().remove(table._element)

    file_path_name = os.path.split(file_path)
    file_name_ext = file_path_name[-1].split('.')
    saved_file_path = file_path_name[0] + '/' + file_name_ext[0] + ' .' + file_name_ext[1]

    if(test):    
        for paragraph in document.paragraphs:

            print(paragraph.text)
            print('=======')
    else:
        document.save(saved_file_path)

    return saved_file_path

if __name__ == '__main__':

    ProcessDocx('./database/File/申辦條件與年費.docx', True)
